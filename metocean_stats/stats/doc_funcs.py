from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.section import WD_ORIENTATION, WD_SECTION, WD_ORIENT
import requests
from io import BytesIO
from datetime import datetime
from docx.oxml import OxmlElement

# Topptekst
def add_header(doc, title, logo_url):
    # Hent logoen fra URL og lagre den midlertidig
    response = requests.get(logo_url)
    logo_bytes = BytesIO(response.content)
    
    # Iterer over hver seksjon i dokumentet
    for section in doc.sections:
        header = section.header
        
        # Legg til en tabell med 1 rad og 2 kolonner
        table = header.add_table(rows=1, cols=2, width=Inches(6))  # Juster bredden etter behov
        table.autofit = True
        
        # Sett kolonnebredder
        for column in table.columns:
            for cell in column.cells:
                cell.width = Inches(3)  # Sett passende bredde for cellene

        # Første celle: Legg til tekst
        cell_text = table.cell(0, 0)
        p = cell_text.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = p.add_run(f"Title: {title}")
        title_run.font.size = Pt(11)

        today = datetime.today().strftime('%d.%m.%Y')
        date_run = p.add_run(f"\nDate: {today}")
        date_run.font.size = Pt(11)

        # Andre celle: Legg til logo
        cell_logo = table.cell(0, 1)
        p_logo = cell_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Høyrejuster logoen
        logo_run = p_logo.add_run()
        logo_run.add_picture(logo_bytes, width=Inches(0.8))  # Juster størrelsen etter behov
        
# Funksjon som legger til et bilde og bildetekst i en tabell
def add_image_with_caption(doc, image_path, caption_text, orientation="portrait"):
    # Lag en ny seksjon om det er landskap
    if orientation == "landscape":
        # Opprett en ny seksjon med landskapsorientering
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.orientation = WD_ORIENTATION.LANDSCAPE
        new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

    # Juster bildebredden slik at det passer til sidebredden
    max_image_width = doc.sections[-1].page_width.inches - doc.sections[-1].left_margin.inches - doc.sections[-1].right_margin.inches

    # Lag en tabell med 2 rader og 1 kolonne
    table = doc.add_table(rows=2, cols=1)
    table.style = 'Table Grid'  # Sett tabellstilen

    # Første celle: bilde
    cell_image = table.cell(0, 0)
    paragraph_image = cell_image.paragraphs[0]
    paragraph_image.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Midtstill paragrafen som inneholder bildet
    run_image = paragraph_image.add_run()

    # Legg til bildet i den første cellen
    image_width = Inches(max_image_width * 0.95) if orientation == "landscape" else Inches(max_image_width)
    run_image.add_picture(image_path, width=image_width)

    # Andre celle: bildetekst
    cell_caption = table.cell(1, 0)
    cell_caption.text = caption_text

    # Sentraliser bildeteksten
    paragraph_caption = cell_caption.paragraphs[0]
    paragraph_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Sentralisert

    # Fjerne plass før og etter avsnitt
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.space_after = 0  # Fjerner eventuell plass etter avsnitt
                paragraph.space_before = 0  # Fjerner eventuell plass før avsnitt

    # Sikre at tabellen ikke har ekstra mellomrom
    table.autofit = True  # Sett til True for å tilpasse størrelsen basert på innhold
    cell_image.width = Inches(max_image_width)
    cell_caption.width = Inches(max_image_width)

    # Legg til en ny seksjon i portrettorientering hvis nødvendig
    if orientation == "landscape":
        portrait_section = doc.add_section(WD_SECTION.NEW_PAGE)
        portrait_section.orientation = WD_ORIENTATION.PORTRAIT
        portrait_section.page_width, portrait_section.page_height = portrait_section.page_height, portrait_section.page_width
        
        
        
        
#-----------Funkjson for Tabell---------------
def add_table_to_doc(doc, df, col_width=50, row_height=0.7, header_text="", header_color='D3D3D3', data_color=None):
    
    if df.index.name or df.index.names != [None]:  
        df = df.reset_index()  # Dette flytter indeksen til en kolonne hvis indeksen har et navn
    
    num_columns = len(df.columns) # Legg til én ekstra kolonne til venstre
    landscape_section = False
    
    # Opprett en ny seksjon i landskapmodus hvis det er mange kolonner
    if num_columns > 10:
        landscape_section = True
        current_section = doc.add_section(WD_SECTION.NEW_PAGE)
        current_section.orientation = WD_ORIENTATION.LANDSCAPE
        current_section.page_width, current_section.page_height = current_section.page_height, current_section.page_width
        current_section.top_margin = Cm(1.5)
        current_section.bottom_margin = Cm(1.5)

    # Legg til hovedoverskriften i dokumentet
    if header_text:
        header_paragraph = doc.add_heading(header_text, level=3)
        header_paragraph.paragraph_format.keep_with_next = True

    # Beregn tilgjengelig bredde for tabellen
    section = doc.sections[-1]
    available_width = section.page_width - section.left_margin - section.right_margin
    column_width = available_width / num_columns

    # Opprett tabellen med riktig antall kolonner, men bare én header-rad
    table = doc.add_table(rows=1, cols=num_columns)
    table.style = 'Table Grid'

    # Fyll header-celler med kolonnenavnene fra DataFrame
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = column
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        cell_pr = hdr_cells[i]._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', header_color)
        cell_pr.append(shd)

    # Fyll de øvrige header-cellene med kolonnenavnene fra DataFrame
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = column
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        cell_pr = hdr_cells[i]._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', header_color)
        cell_pr.append(shd)

    # Legg til data-rader og fyll venstre kolonne med indeksverdiene
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(index)  # Fyll første celle med indeks
        run = row_cells[0].paragraphs[0].runs[0]
        run.font.size = Pt(9)
        if data_color:
            cell_pr = row_cells[0]._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', data_color)
            cell_pr.append(shd)

        # Fyll resten av radene med verdier fra DataFrame
        for j, value in enumerate(row):
            if isinstance(value, (float, int)):
                value = f"{value:.1f}"
            row_cells[j].text = str(value)
            run = row_cells[j].paragraphs[0].runs[0]
            run.font.size = Pt(9)

    # Juster høyden på alle rader
    row_height_cm = Cm(row_height)
    height_twips = int(row_height_cm.pt * 20)
    for row in table.rows:
        tr = row._element
        tr_pr = tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        cant_split.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'true')
        tr_pr.append(cant_split)

        height = OxmlElement('w:trHeight')
        height.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(height_twips))
        tr_pr.append(height)

    # Bytt tilbake til portrettmodus hvis landskap ble brukt
    if landscape_section:
        portrait_section = doc.add_section(WD_SECTION.NEW_PAGE)
        portrait_section.orientation = WD_ORIENTATION.PORTRAIT
        portrait_section.page_width, portrait_section.page_height = portrait_section.page_height, portrait_section.page_width
        portrait_section.left_margin = Cm(2.5)
        portrait_section.right_margin = Cm(2.5)
        portrait_section.top_margin = Cm(1.5)
        portrait_section.bottom_margin = Cm(1.5)




# Function to add superscript text
def add_superscript(paragraph, base_text, superscript_text):
    run = paragraph.add_run(base_text)
    superscript_run = paragraph.add_run(superscript_text)
    superscript_run.font.superscript = True

