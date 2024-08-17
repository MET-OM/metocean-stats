from metocean_stats import plots, tables, stats, maps
from metocean_stats.stats.aux_funcs import *
from docx import Document
from docx.shared import Inches


# Read test data
ds = readNora10File('../tests/data/NORA_test.txt') 
# Create a new Document
doc = Document()



# Add introductory text
intro_text = "This is an automated metocean report produced by MET Norway."
doc.add_paragraph(intro_text)


# Add a title
doc.add_heading('Metocean Report', level=1)


# Add the map figure
output_file='map.png'
maps.plot_points_on_map(lon=[3.35], lat=[60.40],label=['NORA3'],bathymetry='NORA3',output_file=output_file)
doc.add_heading('Figure 1: The figure shows the NORA3 grid points selected for the analysis', level=2)
doc.add_picture(output_file, width=Inches(5))

output_file='wave_100yrs.png'
maps.plot_extreme_wave_map(return_period=100, product='NORA3', title='100-yr return values Hs (NORA3)', set_extent = [0,30,52,73],output_file=output_file)
doc.add_heading('Figure 2: 100-year return period for wind speed at 10 m and 100 m in the Nordics based on NORA3 (period: 1991-2020) using Generalized Pareto distribution', level=2)
doc.add_picture(output_file, width=Inches(5))

output_file='wind_100yrs.png'
maps.plot_extreme_wind_map(return_period=100, product='NORA3',z=10, title='100-yr return values Wind at 100 m (NORA3)', set_extent = [0,30,52,73], output_file=output_file)
doc.add_heading('Figure 3: 100-year return period for significant wave height in the Nordics based on NORA3 (period: 1991-2020) using Gumbel distribution', level=2)
doc.add_picture(output_file, width=Inches(5))



# Add the first table
df = tables.table_monthly_non_exceedance(ds,var= 'W10',step_var=2,output_file=None)
doc.add_heading('Table 1: Monthly non-exceedance', level=2)
table1 = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
table1.style = 'Table Grid'

# Add the header row for the first table
hdr_cells = table1.rows[0].cells
for i, column in enumerate(df.columns):
    hdr_cells[i].text = column

# Add the data rows for the first table
for i, row in df.iterrows():
    row_cells = table1.add_row().cells
    for j, value in enumerate(row):
        row_cells[j].text = str(value)



# Save the document
doc.save('metocean-report.docx')

print("Document created successfully.")
