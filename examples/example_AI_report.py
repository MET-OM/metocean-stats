from metocean_stats import plots, tables, stats, maps
from metocean_stats.stats.aux_funcs import *
from metocean_stats.stats.doc_funcs import *
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
from docx.oxml.ns import qn  # Import the qn function for XML namespaces
from io import BytesIO
from pathlib import Path
import warnings
warnings.filterwarnings("ignore")
import google.generativeai as genai
genai.configure(api_key=os.environ["GEMINI_API_KEY"])
model = genai.GenerativeModel(model_name="gemini-1.5-flash")

##############PROVIDE INFO BY THE USER#################
# Impot data (e.g., use metocean-api to download metocean data)
## For NORA10 data use:
df = readNora10File('../tests/data/NORA_test.txt') 
## For NORA3 data use:
#import pandas as pd
#df = pd.read_csv('../path/to/NORA3.csv', comment="#", index_col=0, parse_dates=True)

# Define names for each variable in the dataframe (df):
var_wind_dir = 'D10' # for wind direction
var_wind = 'W10' # for wind speed
var_hs = 'HS' # for significant wave height
var_wave_dir= 'DIRM' # Mean wave direction
var_tp = 'TP'  # Peak Wave Period
output_folder = 'output_AI_report' # folder where output report and figures will be saved 

# Manually Define varibles for the text in the report
LocationX = "North Sea" # Location name
lon = 3 # Latitude
lat = 60 # Longitude
label = ['NORA'] # Dataset used
chapter = ["wind", "waves"] # Hvilke del/kap som skal være med i rapporten
water_depth = 100
######################################################


# Check if the output directory exists, if not, create it
folder = Path(__file__).parent / output_folder 
if not folder.exists():
    folder.mkdir(parents=True)
folder = str(folder)
# Create a new Document
doc = Document()


starttime = str(df.index[0])
endtime = str(df.index[-1])
# _________________FRONTPAGE_______________

#_____________________Bunntekst og sidetall______

# Legge til en seksjon for å ha forskjellige bunntekster på første side
section = doc.sections[0]
section.different_first_page_header_footer = True

# Sjekke om 'pgNumType' allerede eksisterer, hvis ikke, opprett den
sectPr = section._sectPr  # Får tak i seksjonens XML-elementer
pgNumType = sectPr.xpath('./w:pgNumType')  # Sjekker om pgNumType allerede er der

if not pgNumType:
    # Oppretter pgNumType hvis det ikke eksisterer
    pgNumType = OxmlElement('w:pgNumType')
    sectPr.append(pgNumType)
else:
    pgNumType = pgNumType[0]  # Henter pgNumType-elementet

# Legge til en seksjon for å ha forskjellige bunntekster på første side
section = doc.sections[0]
section.different_first_page_header_footer = True

# Nå setter vi sidetallene til å begynne fra 1 på side 2
section_start = section._sectPr.xpath('./w:pgNumType')[0]
section_start.set(qn('w:start'), "0")

# For å legge til en bunntekst med sidetall på den andre siden
footer = section.footer
footer_paragraph = footer.paragraphs[0]

# Få tilgang til bunnteksten på første og andre side
first_page_footer = section.first_page_footer
main_footer = section.footer

# Slett eksisterende sidetall fra første side (hvis noen)
for paragraph in first_page_footer.paragraphs:
    paragraph.clear()

# Midtstill paragrafen i bunnteksten
footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Her setter vi opp en felttagg for å vise sidetall, og at det starter på 1 fra side 2
page_num_run = footer_paragraph.add_run()
fldChar1 = OxmlElement('w:fldChar')  # Opprettelse av en felttagg
fldChar1.set(qn('w:fldCharType'), 'begin')

instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = "PAGE \\* MERGEFORMAT"

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')

page_num_run._r.append(fldChar1)
page_num_run._r.append(instrText)
page_num_run._r.append(fldChar2)

# Legge til " av " mellom sidetallene
footer_paragraph.add_run(" av ")

# Her setter vi opp en felttagg for å vise totalt antall sider (NUMPAGES)
total_pages_run = footer_paragraph.add_run()
fldChar1 = OxmlElement('w:fldChar')  # Opprettelse av en felttagg
fldChar1.set(qn('w:fldCharType'), 'begin')

instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')

# Bruker en Word-formel til å trekke 1 fra totalt antall sider
instrText.text = "NUMPAGES - 1\\* MERGEFORMAT "

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')

total_pages_run._r.append(fldChar1)
total_pages_run._r.append(instrText)
total_pages_run._r.append(fldChar2)


#-----------TOPPTEKST------------------------

# URL til bildet
logo_url = "https://github.com/MET-OM/metocean-stats/raw/main/docs/files/logo.png"
# Last ned bildet og lagre det midlertidig lokalt
logo_path = folder + "/" +"local_logo.png"
response = requests.get(logo_url)
with open(logo_path, "wb") as file:
    file.write(response.content)



# Legg til topptekst med tittel
title = LocationX+" Metocean Design Basis"
add_header(doc, title, logo_url)

doc.add_paragraph()  # Add a blank line

# Add the document title
title = doc.add_heading(LocationX +" Metocean Design Basis", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.size = Pt(24)  # Increase font size to 24pt

doc.add_paragraph() 

# Add lon and lat
lon_lat_paragraph = doc.add_paragraph(f'Longitude: {lon} \n Latitude: {lat}')
lon_lat_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
lon_lat_run = lon_lat_paragraph.runs[0]
lon_lat_run.font.size = Pt(16)  # Set the date font size to 16pt

doc.add_paragraph()

# Get today's date and format it
today = datetime.today().strftime('%d.%m.%Y')

# Add the date
date_paragraph = doc.add_paragraph(f'Date: {today}')
date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_paragraph.runs[0]
date_run.font.size = Pt(16)  # Set the date font size to 16pt

doc.add_paragraph()

# Add the created by line
created_by_paragraph = doc.add_paragraph('Created by: metocean-stats')
created_by_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
created_by_run = created_by_paragraph.runs[0]
created_by_run.font.size = Pt(16)  # Set the "created by" font size to 16pt

# Download logo and add to the front page
response = requests.get(logo_url)
logo_bytes = BytesIO(response.content)
doc.add_paragraph()  # Add blank line
p_logo = doc.add_paragraph()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_run = p_logo.add_run()
logo_run.add_picture(logo_bytes, width=Inches(3))  # Adjust width as needed

# Add a page break to move to the next page
doc.add_page_break()

#___INNHOLDSFORTEGNELSE____________________________________________

# Add the Table og Contents line
table_of_contents_paragraph = doc.add_paragraph('Table of Contents')
table_of_contents_run = table_of_contents_paragraph.runs[0]
table_of_contents_run.font.size = Pt(24)
# Create a manual Table of Contents (ToC)
toc_paragraph = doc.add_paragraph()

toc_paragraph.add_run("1 INTRODUCTION").bold = True

# Legg til punkt 1.1 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("1.1 Scope")
toc_paragraph.paragraph_format.left_indent = Cm(1)  # Innrykk på 1 cm

# Legg til punkt 1.2 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("1.2 Reference Documents and Links")
toc_paragraph.paragraph_format.left_indent = Cm(1)  # Innrykk på 1 cm

# Legg til hovedoverskrift for 2 ROLES & RESPONSIBILITIES
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("2 ROLES & RESPONSIBILITIES").bold = True

# Legg til hovedoverskrift for 3 DESCRIPTION
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("3 DESCRIPTION").bold = True

# Legg til punkt 3.1 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run(f"3.1 {LocationX} location")
toc_paragraph.paragraph_format.left_indent = Cm(1)  # Innrykk på 1 cm

# Legg til hovedoverskrift for 4 METEOROLOGICAL DATA
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("4 METEOROLOGICAL DATA").bold = True

# Legg til punkt 4.1 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("4.1 Wind")
toc_paragraph.paragraph_format.left_indent = Cm(1)  # Innrykk på 1 cm
# Legg til punkt 4.1.1 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("4.1.1 Wind data")
toc_paragraph.paragraph_format.left_indent = Cm(2)  # Innrykk på 2 cm
# Legg til punkt 4.1.2 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("4.1.2 Wind data analysis")
toc_paragraph.paragraph_format.left_indent = Cm(2)  # Innrykk på 2 cm
# Legg til punkt 4.1.3 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("4.1.3 Long-term wind statistics")
toc_paragraph.paragraph_format.left_indent = Cm(2)  # Innrykk på 2 cm
# Legg til punkt 4.1.4 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("4.1.4 Wind profile")
toc_paragraph.paragraph_format.left_indent = Cm(2)  # Innrykk på 2 cm
# Legg til punkt 4.1.5 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("4.1.5 Wind Gust")
toc_paragraph.paragraph_format.left_indent = Cm(2)  # Innrykk på 2 cm
# Legg til punkt 4.1.6 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("4.1.6 Operational data")
toc_paragraph.paragraph_format.left_indent = Cm(2)  # Innrykk på 2 cm
# Legg til punkt 4.2 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("4.2 Air temperature")
toc_paragraph.paragraph_format.left_indent = Cm(1)  # Innrykk på 1 cm

# Legg til hovedoverskrift for 5 OCEANOGRAPHIC DATA
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("5 OCEANOGRAPHIC DATA").bold = True

# Legg til punkt 5.1 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("5.1 Waves")
toc_paragraph.paragraph_format.left_indent = Cm(1)  # Innrykk på 1 cm
# Legg til punkt 5.1.1 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("5.1.1 Wave data")
toc_paragraph.paragraph_format.left_indent = Cm(2)  # Innrykk på 2 cm
# Legg til punkt 5.1.2 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("5.1.2 Wave data analysis")
toc_paragraph.paragraph_format.left_indent = Cm(2)  # Innrykk på 2 cm
# Legg til punkt 5.1.3 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("5.1.3 Long-term statistics")
toc_paragraph.paragraph_format.left_indent = Cm(2)  # Innrykk på 2 cm
# Legg til punkt 5.1.4 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("5.1.4 Joint distribution of Hs and Tp")
toc_paragraph.paragraph_format.left_indent = Cm(2)  # Innrykk på 2 cm
# Legg til punkt 5.1.5 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("5.1.5 Individual waves and crest heights")
toc_paragraph.paragraph_format.left_indent = Cm(2)  # Innrykk på 2 cm
# Legg til punkt 5.1.6 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("5.1.6 Wave induced seabed currents")
toc_paragraph.paragraph_format.left_indent = Cm(2)  # Innrykk på 2 cm
# Legg til punkt 5.1.7 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("5.1.7 Operational wave analysis")
toc_paragraph.paragraph_format.left_indent = Cm(2)  # Innrykk på 2 cm
# Legg til punkt 5.1.8 med innrykk
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("5.1.8 Wind-wave correlation")
toc_paragraph.paragraph_format.left_indent = Cm(2)  # Innrykk på 2 cm

# Legg til hovedoverskrift for 6 REFERENCES
toc_paragraph = doc.add_paragraph()
toc_paragraph.add_run("6 REFERENCES").bold = True

doc.add_page_break()

# ___________TEXT CONTENT______________





#-----------------------
# 1 INTRODUCTION
doc.add_heading('1 INTRODUCTION', level=1)
doc.add_paragraph(
    f'This metocean report is prepared to support activities related to {LocationX}.'
)

# 1.1 Scope
doc.add_heading('1.1 Scope', level=2)
doc.add_paragraph(
    'This Metocean Design Basis (MDB) outlines metocean statistics and long-term metocean conditions which '
    'are applicable for related activities.'
)

# 1.2 Reference Documents and Links (add placeholder text for now)
doc.add_heading('1.2 Reference Documents and Links', level=2)
doc.add_paragraph('Placeholder for reference documents and relevant links.')

doc.add_paragraph() 

# 2 ROLES & RESPONSIBILITIES
doc.add_heading('2 ROLES & RESPONSIBILITIES', level=1)
doc.add_paragraph(
    'This automated Metocean Design Basis has been produced using the open-source python package '
    'metocean-stats. It is important to note that while metocean-stats provides the tools and methodologies for '
    'generating metocean statistics, the developers of this package do not assume any responsibility or liability for '
    'the accuracy, completeness, or appropriateness of the results generated in this report. The automated results '
    'derived from the use of metocean-stats are solely dependent on the input data and parameters provided by '
    'the user. Therefore, the user is responsible for validating and interpreting the outputs to ensure they meet the '
    'specific requirements of the project.'
)

doc.add_paragraph() 

# 3 DESCRIPTION
doc.add_heading('3 DESCRIPTION', level=1)

# 3.1 {LocationX} location
doc.add_heading(f"3.1 {LocationX} location", level=2)
response = model.generate_content("Write a short geographic description (2 sentences) of the location:"+f'Longitude: {lat}°N,  Latitude: {lon}°E')
location_text = "About the location: "+response.text

doc.add_paragraph(
    f"This report presents meteorological and oceanographic (metocean) data for the {LocationX} field. "
    "Figure 3.1 shows the bathymetry in the area and position of NORA10 [15] for wind, air temperature, and "
    "waves, NORA3 [13,14] The water depth varies from approximately {water_depth} m with reference to mean sea level "
    "(MSL)."
    +location_text
)

doc.add_paragraph() 

#__________________________________________________________
# Add the map figure, figure 3.1

output_file=folder + '/' +'map.png'
maps.plot_points_on_map(lon,lat,label,bathymetry='NORA3',output_file=output_file)
add_image_with_caption(doc, output_file, 'Figure 3.1: The figure shows the NORA3 grid points selected for the analysis', orientation="portrait")
## Add a blank paragraph to ensure there is no extra spacing
doc.add_paragraph()

doc.add_paragraph(
    "Figure 3.2 and Figure 3.3 show the 100-year return values for wind and wave height in the Nordics Seas based "
    "on NORA3 data."
)

doc.add_paragraph()
output_file=folder + '/' +'wind_100yrs.png'
maps.plot_extreme_wind_map(return_period=100, product='NORA3',z=10, title='100-yr return values Wind at 100 m (NORA3)', set_extent = [0,30,52,73], output_file=output_file)
add_image_with_caption(doc, output_file, 'Figure 3.2: 100-year return period for wind speed at 10 m in the Nordics based on NORA3 (period: 1991-2020) using Generalized. Pareto distribution (POT; threshold is the minimum of all annual maxima, method described by [5]).', orientation="portrait")


doc.add_paragraph()
output_file=folder + '/' +'wave_100yrs.png'
maps.plot_extreme_wave_map(return_period=100, product='NORA3', title='100-yr return values Hs (NORA3)', set_extent = [0,30,52,73],output_file=output_file)
add_image_with_caption(doc, output_file, "Figure 3.3: 100-year return period for significant wave height in the Nordics based on NORA3 (period: 1991-2020) using Gumbel distribution (Annual maxima).", orientation="portrait")

doc.add_paragraph()
doc.add_page_break()
#________________________
# 4 METEOROLOGICAL DATA
doc.add_heading('4 METEOROLOGICAL DATA', level=1)

# 4.1 Wind
doc.add_heading(f"4.1 Wind", level=2)
# 4.1.1 Wind data
doc.add_heading(f"4.1.1 Wind data", level=3)
doc.add_paragraph(
    f"Wind data are available from the NORA10 hindcast model operated by the Norwegian Meteorological Institute. "
f"The data cover the period {starttime} - {endtime}. The sample interval is 3 hours. The NORA10 model has "
"a spatial resolution of 10 km. "
f"Wind data from grid point position {lat} N, {lon} E (Figure 3.1) is used in the analysis. "
"The computed wind speed is considered to represent the 1-hour mean wind speed 10 m above sea level."
)

doc.add_paragraph()

# 4.1.2 Wind data analysis
doc.add_heading(f"4.1.2 Wind data analysis", level=3)
doc.add_paragraph(
    f"Figure 4.1 shows the (all-year) wind rose at 10 m above mean sea level for the period {starttime} - {endtime}. The wind "
"rose shows the percentage of observations within each 30° sector. "
"Table 4.1 shows the annual directional sample distribution of non-exceedance of 1-hour mean wind speed. "
"Figure 4.2 shows the directional distribution of mean, P99 and maximum wind speed at 10 m above mean sea "
f"level at the {LocationX} field. "
"Table 4.2 shows the monthly distribution of non-exceedance of 1-hour mean wind speed. "
"Figure 4.3 shows monthly distribution of mean, P99 and maximum 1-hour mean wind speed. "
"Wind roses for each month are displayed in Figure 4.4."
)

doc.add_page_break()

# Legg til Figur 4.1
plots.var_rose(df,var_dir=var_wind_dir,var=var_wind,method='overall',max_perc=20,decimal_places=1, units='Wind speed (m/s)',output_file=folder + "/" +"wind_omni.png")
add_image_with_caption(doc, folder + '/' +'wind_omni.png', "Figure 4.1: All-year wind rose at 10 m above mean sea level for the " + LocationX + " field for the period " + starttime + "-" + endtime + ".", orientation="portrait")

doc.add_page_break()

# Henter ut data til tabell 1
df1 = tables.table_directional_non_exceedance(df, var=var_wind,step_var=2,var_dir=var_wind_dir,output_file=None)
# Legger til tabellen i word
add_table_to_doc(doc, df1, col_width=50, row_height=0.7, header_color='D3D3D3', data_color='D2B48C')


# Legg til Figur 4.2
plots.plot_directional_stats(df,var=var_wind,step_var=0.5, var_dir=var_wind_dir, title = 'W10[m/s]', output_file=folder + "/" +"directional_wind_stats.png")
add_image_with_caption(doc, folder + '/' +'directional_wind_stats.png', "Figure 4.2: Directional distribution of mean, P99 and maximum wind speed at 10 m above mean sea level at the " + LocationX + " field.", orientation="portrait")
doc.add_page_break()

# Hent DataFrame for andre tabell og legger til i word
df2 = tables.table_monthly_non_exceedance(df, var=var_wind, step_var=2, output_file=None)
add_table_to_doc(doc, df2, col_width=50, row_height=0.7, header_color='D3D3D3', data_color='D2B48C')

# Legge til figur 4.3 
plots.plot_monthly_stats(df,var="W10",show=["min","mean","max"],title='Monthly W10 [m/s]',fill_between=["25%","75%"],fill_color_like="mean",output_file=folder + '/' +'monthly_wind_stats.png')
add_image_with_caption(doc, folder + '/' +'monthly_wind_stats.png', "Figure 4.3: Monthly distribution of mean, P99 and maximum wind speed 10 m above mean sea level at the " + LocationX + " field.", orientation="portrait")

# Legg til figur 4.4
plots.var_rose(df,var_dir=var_wind_dir,var=var_wind,method='monthly',max_perc=15,decimal_places=1, units="Wind Speed (m/s)",output_file=folder + "/" +"wind_monthly.png")
add_image_with_caption(doc, folder + '/' +'wind_monthly.png', "Figure 4.4: Monthly wind roses for the " + LocationX + "field for the period " + starttime + "to " + endtime, orientation="landscape")

# 4.1.3 Long-term wind statistics
doc.add_heading(f"4.1.3 Long-term wind statistics", level=3)
doc.add_paragraph(
    f"The long-term distribution of wind speed is modeled in terms of a 3-parameter Weibull distribution. The shape"
"parameter is objectively forced (method of moments) to optimize the fit to the upper tail of the distribution."
f"Figure 4.5 shows the hindcast and fitted distributions of wind speed at the {LocationX} field.")

# Legg til figur 4.5
plots.plot_prob_non_exceedance_fitted_3p_weibull(df,var="W10",output_file=folder + "/" +"prob_non_exceedance_wind_fitted_3p_weibull.png")
add_image_with_caption(doc, folder + '/' +'prob_non_exceedance_wind_fitted_3p_weibull.png', f"Figure 4.5: Hindcast (red) and fitted (blue line) distributions of 1 -hour mean wind speed 10 m [m/s] above sea level at the {LocationX}"
"field.", orientation="portrait")
doc.add_paragraph() 

doc.add_paragraph("Figure 4.6 and Table 4.3 show directional Weibull parameters and corresponding extremes of 1-hour mean "
f"wind speed at the {LocationX} field. The direction extremes are adjusted in agreement with NORSOK Standard"
"N-003:2017.")

# Legg til figur 4.6
plots.plot_directional_return_periods(df,var="W10",units = 'm/s', var_dir="D10",periods = [1,10,100,10000],output_file=folder + "/" +"W10_dir_extremes_Weibull_norsok.png",distribution="Weibull3P")
add_image_with_caption(doc, folder + '/' +'W10_dir_extremes_Weibull_norsok.png', f"Figure 4.6: Directional extreme values of 1 -hour mean wind speed with return period of 1, 10, 100 and 10 000 years, 10 m above sea "
f"level at the {LocationX} field. The direction extremes are adjusted in agreement with NORSOK Standard N - 003:2017", orientation="portrait")
doc.add_paragraph()
doc.add_page_break()

# Hent DataFrame for tabell 4.3
df3= tables.table_directional_return_periods(df,var=var_wind,periods=[1, 10, 100, 10000], units='m/s',var_dir = var_wind_dir,distribution='Weibull3P_MOM', adjustment='NORSOK' ,output_file=None)
add_table_to_doc(doc, df3, col_width=50, row_height=0.7, header_color='D3D3D3', data_color='D2B48C')
doc.add_paragraph()  

doc.add_paragraph("Figure 4.7 and Table 4.4 show monthly Weibull parameters and corresponding extremes.")
doc.add_paragraph()  

# Legg til figur 4.7
plots.plot_monthly_return_periods(df,var="W10",units='m/s', periods=[1,10,100,10000],output_file=folder + "/" +"W10_monthly_extremes.png")
add_image_with_caption(doc, folder + '/' +'W10_monthly_extremes.png', f"Figure 4.7: Monthly extreme values of 1 -hour mean wind speed with return period of 1, 10, 100 and 10 000 years 10 m above sea "
f"level at the {LocationX} field.", orientation="portrait")
doc.add_paragraph()  
doc.add_page_break()

# Legger til tabell 4.4
df4= tables.table_monthly_joint_distribution_Hs_Tp_return_values(df,var_hs=var_hs,var_tp='TP',periods=[1,10,100,10000],output_file=None)
response = model.generate_content("Write 2 sentences about which months have the highest Hs values at the tables:"+str(df4))
Table4_4_text = "Table4.4: "+response.text
add_table_to_doc(doc, df4, col_width=50, row_height=0.7, header_color='D3D3D3', data_color='D2B48C')

doc.add_paragraph(
    f"Table 4.4: "+ Table4_4_text)
doc.add_paragraph()


# 4.1.4 Wind profile
doc.add_heading(f"4.1.4 Wind profile", level=3)
doc.add_paragraph(
    f"Table 4.5 shows the omni-directional extreme values for 1-hour mean wind speed as function of height above "
f"mean sea level (MSL) at the {LocationX} field.")
doc.add_paragraph()  
doc.add_page_break()

# Tabell 4.5
df5 = tables.table_profile_return_values(df,var=[var_wind,'W50','W80','W100','W150'], z=[10, 50, 80, 100, 150], periods=[1, 10, 100, 10000], output_file=None)
response = model.generate_content("Write short summary of the table:"+str(df5))
add_table_to_doc(doc, df5, col_width=50, row_height=0.7, header_color='D3D3D3', data_color='D2B48C')


# 4.1.5 Wind Gust
doc.add_heading(f"4.1.5 Wind Gust", level=3)
doc.add_paragraph(
    f"Table 4.6 shows directional and omni-directional, monthly, and annual extreme values for 10-minute average "
f"wind speed 10 m above mean sea level at the {LocationX} field.")
doc.add_paragraph()  

# MANGLER EN TABELL___________

# 4.1.6 Operational data
doc.add_heading(f"4.1.6 Operational data", level=3)
doc.add_paragraph(
    f"Marine operations may be delayed due to wind speeds exceeding prescribed operational levels (limits) leading "
"to a possible increase in the duration of the operations. Marine operations which must be completed without "
"breaks are called critical. Otherwise, they are termed non-critical. The duration statistics presented in this "
"report is restricted to critical operations only."
"\nFigure 4.8- Figure 4.13 show characteristic durations of operations limited by wind speeds of 10, and 15 m/s for "
"12, 24 and 48 hours. The figures show the expected mean duration and 10-, 50- and 90- percentiles. "
"The figures demonstrate duration characteristics for completing a critical operation including waiting time. "
"Duration [in days] is measured from the day the operation is ready to start. The starting or launching day is "
"assumed to be an arbitrary day within the relevant month.")
doc.add_paragraph() 

# Legg til figur 4.8
# Lager ny variabel slik at man kan endre verdiene i figuren som hentes ut
fig1 =  folder + '/' +'NORA10_monthly_weather_window4_10_12_plot.png'
plots.plot_monthly_weather_window(df,var=var_wind,threshold=10, window_size=12,output_file=fig1)
add_image_with_caption(doc, fig1, f"Figure 4.8: Characteristic durations, including waiting time, to perform operations limited by a wind speed of 10 m/s for 12 hours at "
f"the {LocationX} field.", orientation="portrait")
doc.add_paragraph() 

# Legg til figur 4.9 på samme måte som 4.8 med annderleses verdier
fig2 = folder + '/' +'NORA10_monthly_weather_window4_15_12_plot.png'
plots.plot_monthly_weather_window(df,var=var_wind,threshold=15, window_size=12,output_file= fig2)
add_image_with_caption(doc, fig2, f"Figure 4.9: Characteristic durations, including waiting time, to perform operations limited by a wind speed of 15 m/s for 12 hours at "
f"the {LocationX} field.", orientation="portrait")
doc.add_paragraph() 

# Legg til figur 4.10
fig3 = folder + '/' +'NORA10_monthly_weather_window4_10_24_plot.png'
plots.plot_monthly_weather_window(df,var=var_wind,threshold=10, window_size=24,output_file= fig3)
add_image_with_caption(doc, fig3, f"Figure 4.10: Characteristic durations, including waiting time, to perform operations limited by a wind speed of 10 m/s for 24 hours at "
f"the {LocationX} field.", orientation="portrait")
doc.add_paragraph()  

# Legg til figur 4.11
fig4 = folder + '/' +'NORA10_monthly_weather_window4_15_24_plot.png'
plots.plot_monthly_weather_window(df,var=var_wind,threshold=15, window_size=24,output_file= fig4)
add_image_with_caption(doc, fig4, f"Figure 4.11: Characteristic durations, including waiting time, to perform operations limited by a wind speed of 15 m/s for 24 hours at "
f"the {LocationX} field.", orientation="portrait")
doc.add_paragraph()  

# Legg til figur 4.12
fig5 = folder + '/' +'NORA10_monthly_weather_window4_10_48_plot.png'
plots.plot_monthly_weather_window(df,var=var_wind,threshold=10, window_size=48,output_file= fig5)
add_image_with_caption(doc, fig5, f"Figure 4.12: Characteristic durations, including waiting time, to perform operations limited by a wind speed of 10 m/s for 48 hours at "
f"the {LocationX} field.", orientation="portrait")
doc.add_paragraph() 

# Legg til figur 4.13
fig6 = folder + '/' +'NORA10_monthly_weather_window4_15_48_plot.png'
plots.plot_monthly_weather_window(df,var=var_wind,threshold=15, window_size=48,output_file= fig6)
add_image_with_caption(doc, fig6, f"Figure 4.13: Characteristic durations, including waiting time, to perform operations limited by a wind speed of 15 m/s for 48 hours at "
f"the {LocationX} field.", orientation="portrait")
doc.add_paragraph()  


# 4.2 Air temperature
doc.add_heading(f"4.2 Air temperature", level=2)
doc.add_paragraph(
    f"Figure 4.14 shows the monthly minimum, mean and maximum air temperatures measured at the {LocationX}"
"field."
"\nTable 4.7 shows monthly and annual frequency of non-exceedance of air temperature.")
doc.add_paragraph() 

# Legg til figur 4.14
plots.plot_monthly_stats(df,var="T2m",show=["Maximum","Mean","Minimum"],title = 'Montly T2m [C]' , output_file=folder + "/" +"T2m_monthly_max_mean_min.png")
add_image_with_caption(doc, folder + '/' +'T2m_monthly_max_mean_min.png', f"Figure 4.14: Monthly minimum, mean and maximum air temperature at the {LocationX} field.", orientation="portrait")
doc.add_paragraph() 


doc.add_paragraph(
    f"The extreme air temperatures are estimated with the annual/monthly minimum values. The annual/monthly "
"minimum values are fitted to a Gumbel distribution with maximum likelihood estimation as described in the "
"appendix (Gumbel Distribution)."
"Figure 4.15 and Figure 4.16 shows the monthly distribution of extreme negative and positive air temperatures "
f"for the {LocationX} field. Table 4.8 and Table 4.9 shows the monthly and annual Gumbel parameters and return "
"values for extreme negative air temperature and positive air temperature.")
doc.add_paragraph()  

# Legg til figur 4.15
plots.plot_monthly_return_periods(df,var="T2m",units='C', output_file = folder + "/" +"T2m_monthly_extremes_neg.png",method="minimum",periods=[1,10,100])
add_image_with_caption(doc, folder + '/' +'T2m_monthly_extremes_neg.png', f"Figure 4.15: Monthly distribution of extreme negative air temperature with return period 1, 10 and 100 years (annual probability of "
f"exceedance 0.63, 10 -1 and 10-2) for the {LocationX} field.", orientation="portrait")
doc.add_paragraph()  


# Legg til figur 4.16
plots.plot_monthly_return_periods(df,var="T2m", units ='C', output_file=folder + "/" +"T2m_monthly_extremes_pos.png",method="maximum",periods=[1,10,100])
add_image_with_caption(doc, folder + '/' +'T2m_monthly_extremes_pos.png', f"Figure 4.16: Monthly distribution of extreme positive air temperature with return period 1, 10 and 100 years (annual probability of "
f"exceedance 0.63, 10 -1 and 10-2) for the {LocationX} field.", orientation="portrait")
doc.add_paragraph()  

doc.add_page_break()
# ______________________
# 5 OCEANOGRAPHIC DATA
doc.add_heading('5 OCEANOGRAPHIC DATA', level=1)

# 5.1 Waves
doc.add_heading(f"5.1 Waves", level=2)
# 5.1.1 Wave data
doc.add_heading(f"5.1.1 Wave data", level=3)
doc.add_paragraph(
    f"Wave data for Norwegian coastal waters are available from the NORA10 hindcast model operated by the "
f"Norwegian Meteorological Institute [10]. The data cover the period {starttime} – {endtime}. The sample interval "
"is 3 hours. The NORA10 model has a spatial resolution of 10 km.)"
f"Wave data from grid point position {lat}°N, {lon}°E is used in the analysis. A reference water depth of {water_depth} m "
"below MSL is used.")
doc.add_paragraph() 

# 5.1.2 Wave data analysis
doc.add_heading(f"5.1.2 Wave data analysis", level=3)
doc.add_paragraph(
    "Figure 5.1 shows the all-year wave rose, i.e. the sample direction distribution of significant wave height, at the "
f"{LocationX} field. "
"Table 5.1 shows the direction sample distribution of non-exceedance of significant wave height. "
"Figure 5.3 shows the monthly sample density distribution and basic statistics of significant wave height. "
"Table 5.2 shows the monthly sample distribution of non-exceedance of significant wave height. "
"Wave roses for each month are displayed in Figure 5.4."
)
doc.add_paragraph()  

# Legg til figur 5.1
fig7 = folder + '/' +'wind_omni.png'
plots.var_rose(df,var_dir=var_wind_dir,var=var_wind,method='overall',max_perc=20,decimal_places=1, units='Wave height (m)',output_file=fig7)
add_image_with_caption(doc, fig7, f"Figure 5.1: All-year wave rose for the {LocationX} field for the period {starttime} – {endtime}", orientation="portrait")
doc.add_paragraph()  

# Legg til figur 5.2
fig7a = folder + '/' +'directional_stats.png'
plots.plot_directional_stats(df,var=var_hs,step_var=0.5, var_dir=var_wave_dir, title = 'Hs[m]', output_file=fig7a)
add_image_with_caption(doc, fig7a, f"Figure 5.2: Directional distribution of mean, P99 and maximum of significant wave height at the {LocationX} field.", orientation="portrait")
doc.add_paragraph()  

# MANGLER EN TABELL_________________

# Legg til figur 5.3
fig8 = folder + '/' +'monthly_stats.png'
plots.plot_monthly_stats(df,var=var_hs, title = 'Hs[m]', output_file=fig8)
add_image_with_caption(doc, fig8, f"Figure 5.3: Monthly data distribution (top), mean, P99 and maximum of significant wave height at the {LocationX} field.", orientation="portrait")
doc.add_paragraph() 

# MANGLER EN TABELL_________________

# Legg til figur 5.4
fig9 = folder + '/' +'wind_monthly.png'
plots.var_rose(df,var_dir=var_wind_dir,var=var_wind,method='monthly',max_perc=15,decimal_places=1, units="Wave height (m)",output_file=fig9)
add_image_with_caption(doc, fig9, f"Figure 5.4: Monthly wave roses for the {LocationX} field for the period {starttime} to {endtime}.", orientation="landscape")

# 5.1.3 Long-term statistics
doc.add_heading(f"5.1.3 Long-term statistics", level=3)
doc.add_paragraph(
    "The long-term distribution of significant wave height is modeled in terms of a Weibull distribution. The shape "
"parameter is objectively forced (method of moments) to optimize the fit to the upper tail of the distribution."
f"Figure 5.5 shows the hindcast and fitted distributions of significant wave height at the {LocationX} field."
)
doc.add_paragraph()

# Legg til figur 5.5
plots.plot_prob_non_exceedance_fitted_3p_weibull(df,var="HS",output_file=folder + "/" +"prob_non_exceedance_Hs_fitted_3p_weibull.png")
add_image_with_caption(doc, folder + "/" +"prob_non_exceedance_Hs_fitted_3p_weibull.png", f"Figure 5.5: Hindcast (red dots) and fitted (blue line) distributions of significant wave height at the {LocationX} field.", orientation="portrait")
doc.add_paragraph() 

doc.add_paragraph(
    "Figure 5.6 and Table 5.3 show directional Weibull parameters and corresponding extremes of significant wave "
f"height at the {LocationX} field. The direction extremes are adjusted in agreement with NORSOK Standard"
"N-003:2017."
"Figure 5.7 and Table 5.4 shows the monthly and annual Weibull parameters and corresponding extremes."
)
doc.add_paragraph()

# Legg til figur 5.6
plots.plot_directional_return_periods(df,var="HS",var_dir="DIRP",distribution="Weibull3P",output_file=folder + "/" +"dir_extremes_Weibull_norsok.png")
add_image_with_caption(doc, folder + '/' +'dir_extremes_Weibull_norsok.png', f"Figure 5.6: Adjusted directional extreme values of significant wave height with return period 1, 10, 100 and 10 000 years at the "
f"{LocationX} field. The direction extremes are adjusted in agreement with NORSOK Standard N -003:2017.", orientation="portrait")
doc.add_paragraph()

# MANGLER 1 TABELL___________

# Legg til figur 5.7
plots.plot_monthly_return_periods(df,var="HS",periods=[1,10,100,10000],distribution="Weibull3P_MOM",output_file=folder + "/" +"HS_monthly_extremes.png")
add_image_with_caption(doc, folder + '/' +'HS_monthly_extremes.png', f"Figure 5.7: Monthly extreme values of significant wave height with return period of 1, 10, 100, and 10 000 years.", orientation="portrait")
doc.add_paragraph()  

# MANGLER 1 TABELL___________

# 5.1.4 Joint distribution of Hs and Tp
doc.add_heading(f"5.1.4 Joint distribution of Hs and Tp", level=3)
doc.add_paragraph(
    "A short-term sea state is for most practical purposes reasonably well characterized by the significant wave "
"height, Hs, and the spectral peak period, Tp."
"Table 5.5 shows the annual omni-directional scatter table of Hs and Tp for 65 years. The scatter table is "
"obtained from the 65-year NORA10 hindcast data, and the interval is 3 hours."
"The conditional distribution of spectral peak period (Tp) given significant wave height (Hs) is modeled by a lognormal "
"distribution, as described in the Appendix (Log-normal distribution)."
"Table 5.6 shows the parameters in the log-normal distribution of Tp given HS."
"Figure 5.8 and Table 5.7 show spectral peak period as a function of significant wave height."
"Table 5.8 shows omni-directional extreme significant wave heights and associated spectral peak period."
"Table 5.9 and Table 5.10 show directional and monthly extreme significant wave heights and associated "
"spectral peak period.")
doc.add_paragraph() 


# Legg til figur 5.8
plots.plot_scatter_diagram(df,var1="hs",step_var1=0.5,var2="tp",step_var2=1,output_file=folder + "/" +"tp_for_given_hs.png")
add_image_with_caption(doc, folder + '/' +'tp_for_given_hs.png', f"Figure 5.8: Spectral peak period (Tp) for given significant wave height (Hs) at the {LocationX} field. Heat colormap indicates the density "
"of observations", orientation="portrait")
doc.add_paragraph() 

# MANGLER 4 TABELLER___________

doc.add_paragraph(
    "The LoNoWe model is used to simulate the long-term joint distribution of Hs and Tp, the conditional "
"distribution of spectral peak period (Tp) given significant wave height (Hs). This LoNoWe model is described in"
"Appendix (LoNoWe - Joint distribution of Hs and Tp)."
"Figure 5.10 and Table 5.12 shows the omni-directional contour lines of Hs - Tp for return periods of 1, 10, 100"
"and 10 000 years with the steepness / wave breaking criterion applied.")
doc.add_paragraph()  

# Legg til figur 5.10
plots.plot_joint_distribution_Hs_Tp(
    df,
    var_hs=var_hs,
    var_tp='TP',
    periods=[1,10,100,1000],
    title='Hs-Tp joint distribution',
    output_file=folder + '/' +'Hs.Tp.joint.distribution.png',
    density_plot=True)
add_image_with_caption(doc, folder + '/' +'Hs.Tp.joint.distribution.png', f"Figure 5.10: Contour lines of Hs - Tp including wave breaking criteria/steepness , with return period 1, 10, 100 and 1000 years for "
f"omni-directional waves at the {LocationX} field. Duration of sea state is 3 hours.", orientation="portrait")
doc.add_paragraph() 

# 5.1.5 Individual waves and crest heights
doc.add_heading(f"5.1.5 Individual waves and crest heights", level=3)
doc.add_paragraph(
    "Table 5.13 shows the estimated design wave heights. Extreme value estimates for individual wave heights and "
"wave crests are computed using the Forristal distribution. The wave periods, T_Hmax, are computed from "
"T_Hmax = 0.90 ∙ Tp, where Tp is given in Table 5.8."
"Extreme individual wave heights versus direction sectors are given in Table 5.14. These wave heights are "
"determined from the significant wave heights given in Table 5.9 by assuming that Hmax/Hs for each sector is "
"equal to Hmax/Hs for omni-directional seas and reflect the same relative severity as shown by that table.")
doc.add_paragraph() 


# 5.1.6 Wave induced seabed currents
doc.add_heading(f"5.1.6 Wave induced seabed currents", level=3)
doc.add_paragraph(
    "Table 5.16 and Table 5.17 show wave-induced significant orbital velocity Us and corresponding zero-crossing "
"period Tu, based on JONSWAP and Torsethaugen spectra. The significant wave height and spectral peak data "
"are as given in Table 5.7."
"When the associated spectral peak period is larger than the mean period given in Table 5.14, the most "
"unfavorable orbital velocity from the JONSWAP spectrum should be applied.")
doc.add_paragraph() 

# 5.1.7 Operational wave analysis
doc.add_heading(f"5.1.7 Operational wave analysis", level=3)
doc.add_paragraph(
    "Marine operations may be delayed due to waves exceeding prescribed operational levels (limits) leading to a "
"possible increase in the duration of the operations. Marine operations which must be completed without break "
"are called critical. Otherwise, they are termed non-critical. The duration statistics presented in this report is "
"restricted to critical operations only."
"Figure 5.11 to Figure 5.24 show characteristic durations of operations limited by significant wave heights of 2.0, "
"3.0 and 4.0, 4.5 and 5.5m for 6 hours, and 2.0, 3.0 and 4.0 m for 12, 24 and 48 hours. The figures show the "
"expected mean duration and 10, 50 and 90 percentiles."
"The figures show duration characteristics for completing a critical operation including waiting time. Duration is "
"measured from the day the operation is ready to start. The starting day is considered to be an arbitrary day "
"within the relevant month.")
doc.add_paragraph()  

# Legg til figur 5.11
fig10 = folder + '/' +'NORA10_monthly_weather_window4_12_plot.png'
plots.plot_monthly_weather_window(df,var=var_hs,threshold=2, window_size=6,output_file= fig10)
add_image_with_caption(doc, fig10, f"Figure 5.11: Characteristic durations, including waiting time, to perform operations limited by a significant wave height (Hs) of 2.0 m "
f"for 6 hours at the {LocationX} field.", orientation="portrait")
doc.add_paragraph() 

# Legg til figur 5.12
fig11 = folder + '/' +'NORA10_monthly_weather_window4_12_plot.png'
plots.plot_monthly_weather_window(df,var=var_hs,threshold=3, window_size=6,output_file= fig11)
add_image_with_caption(doc, fig11, f"Figure 5.12: Characteristic durations, including waiting time, to perform operations limited by a significant wave height (Hs) of 3.0 m "
f"for 6 hours at the {LocationX} field.", orientation="portrait")
doc.add_paragraph()  

# Legg til figur 5.13
fig12 = folder + '/' +'NORA10_monthly_weather_window4_12_plot.png'
plots.plot_monthly_weather_window(df,var=var_hs,threshold=4, window_size=6,output_file= fig12)
add_image_with_caption(doc, fig12, f"Figure 5.13: Characteristic durations, including waiting time, to perform operations limited by a significant wave height (Hs) of 4.0 m "
f"for 6 hours at the {LocationX} field.", orientation="portrait")
doc.add_paragraph()  

# Legg til figur 5.14
fig13 = folder + '/' +'NORA10_monthly_weather_window4_12_plot.png'
plots.plot_monthly_weather_window(df,var=var_hs,threshold=4.5, window_size=6,output_file= fig13)
add_image_with_caption(doc, fig13, f"Figure 5.14: Characteristic durations, including waiting time, to perform operations limited by a significant wave height (Hs) of 4.5 m "
f"for 6 hours at the {LocationX} field.", orientation="portrait")
doc.add_paragraph()  

# Legg til figur 5.15
fig14 = folder + '/' +'NORA10_monthly_weather_window4_12_plot.png'
plots.plot_monthly_weather_window(df,var=var_hs,threshold=5.5, window_size=6,output_file= fig14)
add_image_with_caption(doc, fig14, f"Figure 5.15: Characteristic durations, including waiting time, to perform operations limited by a significant wave height (Hs) of 5.5 m "
f"for 6 hours at the {LocationX} field.", orientation="portrait")
doc.add_paragraph()  

# Legg til figur 5.16
fig15 = folder + '/' +'NORA10_monthly_weather_window4_12_plot.png'
plots.plot_monthly_weather_window(df,var=var_hs,threshold=2, window_size=12,output_file= fig15)
add_image_with_caption(doc, fig15, f"Figure 5.16: Characteristic durations, including waiting time, to perform operations limited by a significant wave height (Hs) of 2.0 m "
f"for 12 hours at the {LocationX} field.", orientation="portrait")
doc.add_paragraph() 

# Legg til figur 5.17
fig16 = folder + '/' +'NORA10_monthly_weather_window4_12_plot.png'
plots.plot_monthly_weather_window(df,var=var_hs,threshold=3, window_size=12,output_file= fig16)
add_image_with_caption(doc, fig16, f"Figure 5.17: Characteristic durations, including waiting time, to perform operations limited by a significant wave height (Hs) of 3.0 m "
f"for 12 hours at the {LocationX} field.", orientation="portrait")
doc.add_paragraph()  

# Legg til figur 5.18
fig17 = folder + '/' +'NORA10_monthly_weather_window4_12_plot.png'
plots.plot_monthly_weather_window(df,var=var_hs,threshold=4, window_size=12,output_file= fig17)
add_image_with_caption(doc, fig17, f"Figure 5.18: Characteristic durations, including waiting time, to perform operations limited by a significant wave height (Hs) of 4.0 m "
f"for 12 hours at the {LocationX} field.", orientation="portrait")
doc.add_paragraph()  

# Legg til figur 5.19
fig18 = folder + '/' +'NORA10_monthly_weather_window4_12_plot.png'
plots.plot_monthly_weather_window(df,var=var_hs,threshold=2, window_size=24,output_file= fig18)
add_image_with_caption(doc, fig18, f"Figure 5.19: Characteristic durations, including waiting time, to perform operations limited by a significant wave height (Hs) of 2.0 m "
f"for 24 hours at the {LocationX} field.", orientation="portrait")
doc.add_paragraph()  

# Legg til figur 5.20
fig19 = folder + '/' +'NORA10_monthly_weather_window4_12_plot.png'
plots.plot_monthly_weather_window(df,var=var_hs,threshold=3, window_size=24,output_file= fig19)
add_image_with_caption(doc, fig19, f"Figure 5.20: Characteristic durations, including waiting time, to perform operations limited by a significant wave height (Hs) of 3.0 m "
f"for 24 hours at the {LocationX} field.", orientation="portrait")
doc.add_paragraph() 

# Legg til figur 5.21
fig20 = folder + '/' +'NORA10_monthly_weather_window4_12_plot.png'
plots.plot_monthly_weather_window(df,var=var_hs,threshold=4, window_size=24,output_file= fig20)
add_image_with_caption(doc, fig20, f"Figure 5.21: Characteristic durations, including waiting time, to perform operations limited by a significant wave height (Hs) of 4.0 m "
f"for 24 hours at the {LocationX} field.", orientation="portrait")
doc.add_paragraph() 

# Legg til figur 5.22
fig21 = folder + '/' +'NORA10_monthly_weather_window4_12_plot.png'
plots.plot_monthly_weather_window(df,var=var_hs,threshold=2, window_size=48,output_file= fig21)
add_image_with_caption(doc, fig21, f"Figure 5.22: Characteristic durations, including waiting time, to perform operations limited by a significant wave height (Hs) of 2.0 m "
f"for 48 hours at the {LocationX} field.", orientation="portrait")
doc.add_paragraph() 

# Legg til figur 5.23
fig22 = folder + '/' +'NORA10_monthly_weather_window4_12_plot.png'
plots.plot_monthly_weather_window(df,var=var_hs,threshold=3, window_size=48,output_file= fig22)
add_image_with_caption(doc, fig22, f"Figure 5.23: Characteristic durations, including waiting time, to perform operations limited by a significant wave height (Hs) of 3.0 m "
f"for 48 hours at the {LocationX} field.", orientation="portrait")
doc.add_paragraph() 

# Legg til figur 5.24
fig23 = folder + '/' +'NORA10_monthly_weather_window4_12_plot.png'
plots.plot_monthly_weather_window(df,var=var_hs,threshold=4, window_size=48,output_file= fig23)
add_image_with_caption(doc, fig23, f"Figure 5.24: Characteristic durations, including waiting time, to perform operations limited by a significant wave height (Hs) of 4.0 m "
f"for 48 hours at the {LocationX} field.", orientation="portrait")
doc.add_paragraph() 
#_______________________________________


# 5.1.8 Wind-wave correlation
doc.add_heading(f"5.1.8 Wind-wave correlation", level=3)
doc.add_paragraph(
    "Significant wave height (Hs) and wind speed (U) at the Luna field based on the NORA10 hindcast data, from "
f"{starttime} to {endtime}."
"The relationship between significant wave height for a given wind speed, denoted HS(U), is modeled by the "
"power functions for mean (μ) and standard deviation (σ):")
# "μ(𝐻𝑠(𝑈)) = 𝑎 + 𝑏 𝑈^(𝑐+𝑑.𝑈)"
# "σ(𝐻𝑠(𝑈)) = 𝑎 + 𝑏 𝑈^(𝑐+𝑑.𝑈)"
# Function to add superscript text

# Adding the mean (μ) formula with superscripts
p = doc.add_paragraph("μ(Hs(U)) = a + b U")
add_superscript(p, '', '(c + d.U)')

# Adding the standard deviation (σ) formula with superscripts
p = doc.add_paragraph("σ(Hs(U)) = a + b U")
add_superscript(p, '', '(c + d.U)')

doc.add_paragraph("\nwhere the coefficients a, b, c and d are determined from the data and given in Table 5.18.")
doc.add_paragraph() 

# MANGLER EN TABELL___________

doc.add_paragraph(
    r"An approximation to the 90% confidence interval is obtained under the assumption of a Gaussian distribution:"
r"𝐻𝑆(𝑈)5% = μ(𝐻𝑠(𝑈)) − 1. 65 σ 𝐻𝑠( (𝑈))"
r"\n𝐻𝑆(𝑈)5% = μ 𝐻𝑠( (𝑈)) + 1. 65 σ 𝐻𝑠( (𝑈))"
r"Table 5.19 shows the scatter table of significant wave height and wind speed."
r"Figure 5.25 shows the scatter plot of Hs and U with mean Hs and 90\% confidence interval and the "
r"corresponding modeled Hs overlain. HS(U) is tabulated in Table 5.20."
r"The modeled significant wave height for estimated extreme wind speed return values are given in Table 5.21.")
doc.add_paragraph() 

# MANGLER 1 TABELL___________

# Legg til figur 5.25
plots.plot_hs_for_given_wind(df,var_hs="hs",var_wind="W10",output_file=folder + "/" +"hs_for_given_wind.png")
add_image_with_caption(doc, folder + '/' +'hs_for_given_wind.png', f"Figure 5.25: Relationship between significant wave height H S for a given wind speed U at the Luna field; data and model "
"extrapolation.", orientation="portrait")
doc.add_paragraph() 

# MANGLER 2 TABELLER___________

# 6 REFERENCES
doc.add_heading('6 REFERENCES', level=1)
doc.add_paragraph(
    "1. Sanne Muis, Maialen Irazoqui Apecechea, José Antonio Álvarez, Martin Verlaan, Kun Yan, Job Dullaart, "
"Jeroen Aerts, Trang Duong, Rosh Ranasinghe, Dewi le Bars, Rein Haarsma, Malcolm Roberts, (2022): Global sea "
"level change time series from 1950 to 2050 derived from reanalysis and high resolution CMIP6 climate "
"projections. Copernicus Climate Change Service (C3S) Climate Data Store (CDS). DOI: 10.24381/cdf.a6d42d60 "
"(Accessed on 01-July-2024). "
"2. NORSOK Standard N-003:2017, Action and action effects. Revision No. 3, April 2017. URL: "
"http://www.standard.no/ "
"3. ISO 19901-1:2015, Petroleum and natural gas industries - Specific requirements for offshore structures Part "
"1: Metocean design and operating conditions. International Organization for Standardization. URL "
"http://www.standard.no/ "
"4. DNV-RP-C205 Environmental conditions and environmental loads, April 2014. URL: "
"https://rules.dnvgl.com/servicedocuments/dnv "
"5. Outten, Stephen & Sobolowski, Stefan. (2021). Extreme wind projections over Europe from the Euro-CORDEX "
"regional climate models. Weather and Climate Extremes. 33. 100363. 10.1016/j.wace.2021.100363. "
"6. T. Moan, Z. Gao, and E. Ayala-Uraga. Uncertainty of wave-induced response of marine structures due to "
"long-term variation of extratropical wave conditions. Marine structures, 18(4):359–382, 2005. ISSN 0951-8339. "
"doi:10.1016/j.marstruc.2005.11.001. "
"7. G. L. DNV. Recommended practice DNVGL-RP-C205 environmental conditions and environmental loadf. "
"Høvik: DNV GL AS, 2017. "
"8. G. L. DNV. Recommended practice DNVGL-RP-C205 environmental conditions and environmental loadf. "
"Høvik: DNV GL AS, 2019. "
"9. Forristall, G. Z. (1978), On the statistical distribution of wave heights in a storm, Journal of Geophysical "
"Research, Vol. 83, No. C5, pp. 2353-2358. "
"10. Meer, Ruurd & de Boer, M. & Liebich, Viola & Hallers, Cato & Veldhuis, Marcel & Ree, Karin. (2016). Ballast "
"Water Risk Indication for the North Sea. Coastal Management. 44. 1-22. 10.1080/08920753.2016.1233794. "
"11. ISO 19901-1:2015, Petroleum and natural gas industries - Specific requirements for offshore structures – "
"Part 1: Metocean design and operating conditions. International Organization for Standardization. "
"12. Kristensen et al (2024), NORA-Surge: A storm surge hindcast for the Norwegian Sea, the North Sea and the "
"Barents Sea. "
"13. Breivik, Ø., Carrasco, A., Haakenstad, H., Aarnes, O. J., Behrens, A., Bidlot, J.-R., Björkqvist, J.-V., Bohlinger, P., "
"Furevik, B. R., Staneva, J., and Reistad, M.(2022). The impact of a reduced high-wind Charnock parameter on "
"wave growth with application to the North Sea, the Norwegian Sea, and the Arctic Ocean. Journal of "
"Geophysical Research: Oceans, 127, e2021JC018196. https://doi.org/10.1029/2021JC018196 "
"14. Haakenstad, H., Breivik, Ø., Furevik, B., Reistad, M., Bohlinger, P., & Aarnes, O. J. (2021). NORA3: A "
"nonhydrostatic high-resolution hindcast of the North Sea, the Norwegian Sea, and the Barents Sea. Journal of "
"Applied Meteorology and Climatology. https://doi.org/10.1175/JAMC-D-21-0029.1 "
"15. Reistad, M., Ø. Breivik, H. Haakenstad, O. J. Aarnes, B. R. Furevik, and J.-R. Bidlot, 2011: A high-resolution "
"hindcast of wind and waves for the North Sea, the Norwegian Sea, and the Barents Sea. J. Geophys. Res., 116, "
"C05019, https://doi.org/10.1029/2010JC006402. "
)

#----- SAVE AND OPEN ----
# Save the document
output_filename = folder + '/' +'metocean-report'

# if overwrite fails (file exists/permission) write new file
try:
    doc.save(output_filename+".docx")
except:
    counter = 0
    while os.path.exists(output_filename+"_"+str(counter)+".docx"):
        counter += 1
    doc.save(output_filename+"_"+str(counter)+".docx")

# Open the document with the default application
#os.system(output_filename)  # For Windows

print("Document created successfully")
